#!/usr/bin/env python3
"""
PMS User Manual Generator
Generates a .docx User Manual for the RSMS Planned Maintenance System (PMS) module,
with annotated screenshots highlighting key UI areas.
"""

import os
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

ASSETS_DIR = Path("scripts/manual_assets")
ANNOTATED_DIR = ASSETS_DIR / "annotated"
OUTPUT_FILE = "PMS_User_Manual.docx"

ANNOTATED_DIR.mkdir(parents=True, exist_ok=True)

HIGHLIGHT_COLOR = (220, 53, 69)   # red
HIGHLIGHT_WIDTH = 4

def annotate_screenshot(src_filename, dst_filename, boxes, labels=None):
    """
    Load a screenshot JPEG, draw highlight boxes on it, save annotated PNG.
    boxes: list of (x1, y1, x2, y2) tuples in pixel coordinates
    labels: optional list of strings to place inside each box
    """
    src = ASSETS_DIR / src_filename
    dst = ANNOTATED_DIR / dst_filename
    img = Image.open(src).convert("RGB")
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 14)
    except Exception:
        font = ImageFont.load_default()

    for i, box in enumerate(boxes):
        x1, y1, x2, y2 = box
        for t in range(HIGHLIGHT_WIDTH):
            draw.rectangle(
                [x1 - t, y1 - t, x2 + t, y2 + t],
                outline=HIGHLIGHT_COLOR
            )
        if labels and i < len(labels) and labels[i]:
            lbl = labels[i]
            draw.rectangle([x1, y1 - 20, x1 + len(lbl) * 8 + 8, y1], fill=HIGHLIGHT_COLOR)
            draw.text((x1 + 4, y1 - 19), lbl, fill=(255, 255, 255), font=font)

    img.save(str(dst), "PNG")
    return dst


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    return h


def add_screenshot(doc, img_path, caption, width_inches=6.0):
    """Add an annotated screenshot followed by a caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p


def add_step(doc, step_num, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(11)
    return p


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    section = doc.sections[0]
    section.top_margin = Cm(top * 2.54)
    section.bottom_margin = Cm(bottom * 2.54)
    section.left_margin = Cm(left * 2.54)
    section.right_margin = Cm(right * 2.54)


def add_cover_page(doc):
    section = doc.sections[0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n\n")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("RSMS")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Remote Ship Management System")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0x52, 0xBA, 0xF3)

    doc.add_paragraph()

    mod_p = doc.add_paragraph()
    mod_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mod_run = mod_p.add_run("PMS Module — User Manual")
    mod_run.font.size = Pt(28)
    mod_run.font.bold = True
    mod_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    doc.add_paragraph()

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc_p.add_run(
        "Planned Maintenance System\n"
        "Step-by-step guide for vessel crew, superintendents, and fleet managers"
    )
    desc_run.font.size = Pt(13)
    desc_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("\n\n\n\n")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run("April 2026  |  Version 1.0")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def add_toc_placeholder(doc):
    add_heading(doc, "Table of Contents", level=1)
    sections = [
        "1.  Overview & Navigation ................ 3",
        "2.  Dashboard .............................. 4",
        "3.  Components ............................ 6",
        "4.  Work Orders ........................... 8",
        "5.  Running Hours ........................ 11",
        "6.  Spares Inventory .................... 13",
        "7.  Stores Inventory .................... 15",
        "8.  Reports ............................... 17",
        "9.  Modify PMS ........................... 19",
    ]
    for s in sections:
        p = doc.add_paragraph(s)
        p.runs[0].font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()


def build_manual():
    doc = Document()
    set_page_margins(doc)

    add_cover_page(doc)
    add_toc_placeholder(doc)

    # ─────────────────────────────────────────────
    # 1. OVERVIEW & NAVIGATION
    # ─────────────────────────────────────────────
    add_heading(doc, "1.  Overview & Navigation", level=1)
    add_body(doc,
        "The Planned Maintenance System (PMS) is a core module within RSMS designed to help "
        "vessel crews and shore-based fleet managers plan, execute, and monitor all scheduled "
        "and unscheduled maintenance activities across the vessel fleet. The PMS integrates "
        "component management, work order scheduling, running-hour tracking, inventory "
        "management, and reporting into a single unified interface."
    )

    add_heading(doc, "1.1  Accessing PMS", level=2)
    add_body(doc,
        "PMS is accessed through the top navigation bar. Click the PMS tab at the top of "
        "the screen to enter the module. The left-hand sidebar provides access to all PMS "
        "sub-modules."
    )

    img = annotate_screenshot(
        "01_dashboard.jpg", "00_overview_nav.png",
        boxes=[
            (82, 64, 265, 710),   # sidebar
            (450, 5, 545, 58),    # PMS top tab
        ],
        labels=["Sidebar", "PMS Tab"]
    )
    add_screenshot(doc, img,
        "Figure 1.1 – PMS sidebar navigation (highlighted) and PMS top-navigation tab"
    )

    add_heading(doc, "1.2  Sidebar Sub-Modules", level=2)
    add_body(doc, "The sidebar contains the following PMS sub-modules:")
    items = [
        ("Dashboard", "Fleet-wide KPIs, compliance overview, and maintenance trend charts."),
        ("Components", "Hierarchical register of all vessel machinery and equipment."),
        ("Work Orders", "Scheduled, due, overdue, and unplanned maintenance tasks."),
        ("Running Hrs", "Machinery operating-hour tracking and utilization rate."),
        ("Spares", "Spare parts inventory — ROB, criticality, and consumption history."),
        ("Stores", "Consumable stores inventory (lubes, chemicals, provisions, others)."),
        ("Reports", "Generate and download standard maintenance and inventory reports."),
        ("Modify PMS", "Submit and approve change requests to the PMS structure."),
        ("Admin", "System configuration, data masters, and user management."),
    ]
    for name, desc in items:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # 2. DASHBOARD
    # ─────────────────────────────────────────────
    add_heading(doc, "2.  Dashboard", level=1)
    add_body(doc,
        "The Dashboard provides a high-level overview of the vessel's maintenance health. "
        "It is the landing screen when you enter the PMS module and is the primary tool for "
        "superintendents and chief engineers to monitor fleet performance at a glance."
    )

    add_heading(doc, "2.1  Management vs. Operation Views", level=2)
    add_body(doc,
        "The Dashboard has two views toggled by the Management and Operation tabs in the header:"
    )
    add_bullet(doc, "Fleet-level KPIs: overdue work orders, compliance percentages, and PMS change requests. Ideal for superintendents monitoring multiple vessels.", bold_prefix="Management view: ")
    add_bullet(doc, "Day-to-day operational metrics: work order planner, outstanding jobs, and unplanned maintenance tracking. Designed for chief engineers and deck officers.", bold_prefix="Operation view: ")

    img = annotate_screenshot(
        "01_dashboard.jpg", "02_dashboard_tabs.png",
        boxes=[(560, 88, 780, 126)],
        labels=["View Switcher"]
    )
    add_screenshot(doc, img,
        "Figure 2.1 – Management / Operation tab switcher in the Dashboard header"
    )

    add_heading(doc, "2.2  Key Metrics Cards", level=2)
    add_body(doc,
        "Each metric card shows a gauge chart with a percentage and a raw count. Red indicates "
        "critical attention required, amber is a warning level, and green indicates healthy status."
    )
    add_bullet(doc, "Percentage of work orders past their due date across all equipment.")
    add_bullet(doc, "Percentage of overdue work orders specifically for critical equipment items.")
    add_bullet(doc, "Work orders scheduled but not yet executed — a visual breakdown by status.")
    add_bullet(doc, "Total number of postponed work orders and their percentage of total.")
    add_bullet(doc, "Percentage of unplanned (reactive) maintenance versus total maintenance.")

    img = annotate_screenshot(
        "01_dashboard.jpg", "02_dashboard_kpis.png",
        boxes=[
            (120, 155, 400, 710),  # left KPI column
            (415, 155, 960, 450),  # center trend chart
            (970, 155, 1260, 710), # right stock status
        ],
        labels=["KPI Gauges", "6-Month Trend", "Stock Status"]
    )
    add_screenshot(doc, img,
        "Figure 2.2 – Dashboard KPI gauges (left), 6-month maintenance trend chart (center), "
        "and spares stock status (right)"
    )

    add_heading(doc, "2.3  Applying Filters", level=2)
    add_body(doc,
        "Use the Filters button (top-right of the Dashboard) to narrow metrics by vessel, "
        "department (HOD scope), criticality level, or date range. Filters apply across all "
        "KPI cards and trend charts on the active view."
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # 3. COMPONENTS
    # ─────────────────────────────────────────────
    add_heading(doc, "3.  Components", level=1)
    add_body(doc,
        "The Components module is the master register of all machinery and equipment on the "
        "vessel. It is organized as a hierarchical tree following the SFI (Ship Function Index) "
        "classification standard. Each component can have associated maintenance jobs, spare "
        "parts, and technical specifications."
    )

    add_heading(doc, "3.1  Navigating the Component Tree", level=2)
    add_body(doc,
        "The left panel displays the equipment hierarchy. Click the arrow (›) beside any "
        "category to expand it and reveal sub-components. Use the Expand / Collapse buttons "
        "in the panel header to expand or collapse all nodes at once."
    )

    img = annotate_screenshot(
        "02_components.jpg", "03_components_tree.png",
        boxes=[
            (105, 193, 447, 520),  # component tree panel
            (105, 140, 695, 180),  # search & filter bar
        ],
        labels=["Component Tree", "Search & Filter"]
    )
    add_screenshot(doc, img,
        "Figure 3.1 – Component tree (left panel) with SFI category hierarchy, "
        "and search/filter bar above"
    )

    add_heading(doc, "3.2  Searching and Filtering", level=2)
    add_body(doc, "The filter bar at the top of the Components screen provides:")
    add_bullet(doc, "Full-text search by component name, SFI code, or fleet equipment code.")
    add_bullet(doc, "Critical Item filter to show only critical equipment (Critical / Non-Critical / All Items).")
    add_bullet(doc, "Clear button to reset all active filters.")

    add_heading(doc, "3.3  Viewing Component Details", level=2)
    add_body(doc,
        "Click any component in the tree to open its detail panel on the right side of the "
        "screen. The detail panel shows:"
    )
    add_bullet(doc, "Technical specs: Maker, Model, Serial Number, Drawing Number, Location, Fleet Equipment Code.")
    add_bullet(doc, "Associated maintenance jobs with their frequency and last execution date.")
    add_bullet(doc, "Linked spare parts and their current ROB (Remaining on Board) levels.")
    add_bullet(doc, "Full maintenance history for the selected component.")

    add_heading(doc, "3.4  Exporting Components", level=2)
    add_body(doc,
        "Click the Export button (top-right) to download the current component list as an "
        "Excel (.xlsx) file. The export respects any active search or criticality filters."
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # 4. WORK ORDERS
    # ─────────────────────────────────────────────
    add_heading(doc, "4.  Work Orders", level=1)
    add_body(doc,
        "Work Orders (WO) are the core operational tool for vessel crew. Every maintenance "
        "job that is due, scheduled, or executed is managed here. The Work Orders screen is "
        "designed to be the daily workflow tool for Chief Engineers, 2nd Engineers, and "
        "department officers."
    )

    add_heading(doc, "4.1  Status Tabs", level=2)
    add_body(doc, "Work orders are grouped into the following status tabs:")
    wo_tabs = [
        ("Scheduled", "Jobs that have been planned and have a future due date."),
        ("Due", "Jobs that are due within the current period."),
        ("Overdue", "Jobs past their due date that have not been completed."),
        ("Postponed", "Jobs that have been deferred with an approved justification."),
        ("Unplanned", "Ad-hoc or reactive maintenance jobs not in the planned schedule."),
        ("Pending Approval", "Completed jobs awaiting superintendent sign-off."),
        ("Completed", "Fully executed and approved maintenance jobs."),
    ]
    for tab, desc in wo_tabs:
        add_bullet(doc, f" {desc}", bold_prefix=f"{tab}: ")

    img = annotate_screenshot(
        "03_work_orders.jpg", "04_wo_tabs.png",
        boxes=[
            (355, 88, 1035, 126),   # status tabs
            (1035, 88, 1265, 126),  # Planner + Unplanned WO buttons
        ],
        labels=["Status Tabs", "Action Buttons"]
    )
    add_screenshot(doc, img,
        "Figure 4.1 – Work Orders status tabs (Scheduled, Due, Overdue, etc.) "
        "and action buttons in the header"
    )

    add_heading(doc, "4.2  Filters", level=2)
    add_body(doc, "The filter bar below the tabs allows you to narrow the work order list:")
    add_bullet(doc, "Search by work order number, job title, or component name.")
    add_bullet(doc, "Period filter: select a date range to see only work orders due within that window.")
    add_bullet(doc, "All Ranks dropdown: filter by the rank assigned to the job (e.g., Chief Engineer, 2nd Engineer).")
    add_bullet(doc, "Criticality dropdown: show only Critical or Non-Critical equipment jobs.")
    add_bullet(doc, "Clear: reset all active filters.")

    img = annotate_screenshot(
        "03_work_orders.jpg", "04_wo_filters.png",
        boxes=[(105, 140, 1010, 180)],
        labels=["Filter Bar"]
    )
    add_screenshot(doc, img,
        "Figure 4.2 – Work Orders filter bar: search, period, rank, and criticality filters"
    )

    add_heading(doc, "4.3  Scheduling a Work Order", level=2)
    add_body(doc, "To plan/schedule a work order:")
    add_step(doc, 1, "Navigate to the Scheduled or Due tab.")
    add_step(doc, 2, "Click the work order row to open the Work Order detail panel.")
    add_step(doc, 3, "Set the Planned Date and assign a Rank (responsible officer).")
    add_step(doc, 4, "Click Save. The work order is now confirmed as scheduled.")

    add_heading(doc, "4.4  Executing (Reporting) a Work Order", level=2)
    add_body(doc, "To report a completed maintenance job:")
    add_step(doc, 1, "Open the work order and click Execute.")
    add_step(doc, 2, "Fill in the Findings, Actions Taken, and any Measurement readings.")
    add_step(doc, 3, "Record spare parts consumed (linked automatically to Spares ROB).")
    add_step(doc, 4, "Submit the report. The work order moves to Pending Approval status.")

    add_heading(doc, "4.5  Postponing a Work Order", level=2)
    add_body(doc, "To postpone a work order:")
    add_step(doc, 1, "Click on the work order in the Due or Overdue tab.")
    add_step(doc, 2, "Click the Postpone button.")
    add_step(doc, 3, "Select the postponement reason (e.g., Port Stay, No Spares Available, Weather).")
    add_step(doc, 4, "Enter the new target date and submit for approval.")

    add_heading(doc, "4.6  Planner View", level=2)
    add_body(doc,
        "Click the Planner button (calendar icon, top-right) to switch to a calendar-based "
        "view. The Planner shows all work orders across a weekly or monthly timeline, making "
        "it easy to identify periods with heavy maintenance load and redistribute work."
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # 5. RUNNING HOURS
    # ─────────────────────────────────────────────
    add_heading(doc, "5.  Running Hours", level=1)
    add_body(doc,
        "Running Hours (RH) tracks the total operating hours of machinery components. "
        "Many PMS jobs are triggered by running hours rather than calendar dates — for "
        "example, 'overhaul every 4,000 hours'. Accurate RH updates are essential for "
        "the maintenance schedule to function correctly."
    )

    add_heading(doc, "5.1  Overview Screen", level=2)
    add_body(doc,
        "The Overview tab shows all components with running-hour triggers. The table "
        "columns include:"
    )
    add_bullet(doc, "Component Name and Code — identifies the machinery item.")
    add_bullet(doc, "Running Hours — the current total hour counter.")
    add_bullet(doc, "Last Updated — date and time of the most recent RH update.")
    add_bullet(doc, "Utilization Rate (Monthly) — average daily hours calculated over the selected period.")
    add_bullet(doc, "Inherited RH — whether this component inherits its running hours from a parent component.")
    add_bullet(doc, "Updated By — the user who last submitted the running hours.")

    img = annotate_screenshot(
        "04_running_hours.jpg", "05_rh_overview.png",
        boxes=[
            (105, 193, 1258, 250),  # table header row
            (1035, 88, 1265, 126),  # Bulk Update RH + Export buttons
            (580, 88, 775, 126),    # Overview/History tabs
        ],
        labels=["Table Columns", "Bulk Update / Export", "Overview / History"]
    )
    add_screenshot(doc, img,
        "Figure 5.1 – Running Hours overview: table columns, Overview/History tabs, "
        "and Bulk Update / Export buttons"
    )

    add_heading(doc, "5.2  Updating Running Hours", level=2)
    add_body(doc, "There are two ways to update running hours for a single component:")
    add_bullet(doc, "Enter the actual total meter reading (e.g., 12,450 hours). The system calculates the increment automatically.", bold_prefix="Set Total: ")
    add_bullet(doc, "Enter the hours run since the last update (e.g., 120 hours). The system adds this to the existing total.", bold_prefix="Add Delta: ")
    add_body(doc, "To update a single component:")
    add_step(doc, 1, "Click the component row in the Running Hours table.")
    add_step(doc, 2, "In the update panel, choose Set Total or Add Delta.")
    add_step(doc, 3, "Enter the value and click Save.")

    add_heading(doc, "5.3  Bulk Update", level=2)
    add_body(doc,
        "To update running hours for multiple components at once, click the Bulk Update RH "
        "button (top-right). A dialog opens with an editable table of all components — "
        "enter hours for each and submit in one action."
    )

    add_heading(doc, "5.4  Meter Replacement", level=2)
    add_body(doc,
        "When a machinery meter is replaced and reset to zero, use the Meter Replacement "
        "function so the system does not treat the reset as a negative delta. The cumulative "
        "history is preserved while the new meter starts counting from zero."
    )

    add_heading(doc, "5.5  History Tab", level=2)
    add_body(doc,
        "Click the History tab in the header to view a full audit trail of all running-hour "
        "updates for the selected component, including who submitted each entry and when."
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # 6. SPARES INVENTORY
    # ─────────────────────────────────────────────
    add_heading(doc, "6.  Spares Inventory", level=1)
    add_body(doc,
        "The Spares module manages the inventory of spare parts held on board. It provides "
        "real-time ROB (Remaining on Board) figures, criticality information, and consumption "
        "history. Spares are linked to specific components and are automatically consumed "
        "when recorded in a Work Order execution report."
    )

    add_heading(doc, "6.1  Inventory View", level=2)
    add_body(doc,
        "The Inventory tab is the default view. It shows a table of all spare parts for the "
        "selected component (or all components if none selected). Key columns include:"
    )
    add_bullet(doc, "Part Code and Part Name — unique identifier and description of the spare part.")
    add_bullet(doc, "Component — the equipment item this spare belongs to.")
    add_bullet(doc, "Part Number — manufacturer part number for ordering reference.")
    add_bullet(doc, "Criticality — whether the spare is classified as Critical.")
    add_bullet(doc, "ROB (Remaining on Board) — current quantity on board.")
    add_bullet(doc, "Min Stock — minimum acceptable quantity before a low-stock alert is triggered.")

    img = annotate_screenshot(
        "05_spares.jpg", "06_spares_layout.png",
        boxes=[
            (105, 193, 432, 568),   # component search tree
            (436, 193, 1260, 250),  # table header
            (510, 88, 805, 126),    # Inventory/Location/History tabs
            (436, 140, 975, 180),   # filter bar
        ],
        labels=["Component Tree", "Table Header", "View Tabs", "Filters"]
    )
    add_screenshot(doc, img,
        "Figure 6.1 – Spares Inventory: component tree (left), view tabs, filter bar, "
        "and inventory table"
    )

    add_heading(doc, "6.2  Component Tree Navigation", level=2)
    add_body(doc,
        "The left panel mirrors the equipment hierarchy from the Components module. "
        "Click any category or component to filter the spare parts table to show only "
        "parts associated with that branch of the hierarchy."
    )

    add_heading(doc, "6.3  Filters", level=2)
    add_body(doc, "The filter bar provides:")
    add_bullet(doc, "Search by part name or component name.")
    add_bullet(doc, "Criticality dropdown: filter to Critical or Non-Critical spares.")
    add_bullet(doc, "Rotation Item dropdown: filter to rotation (exchangeable) items.")
    add_bullet(doc, "Stock dropdown: show All, In Stock, or Out of Stock items.")

    add_heading(doc, "6.4  Updating ROB", level=2)
    add_body(doc, "To manually update a spare part's ROB:")
    add_step(doc, 1, "Click the part row in the table.")
    add_step(doc, 2, "Edit the ROB quantity field in the detail panel.")
    add_step(doc, 3, "Click Save. The change is recorded in the History tab.")

    add_heading(doc, "6.5  Bulk Update Spares", level=2)
    add_body(doc,
        "Click Bulk Update Spares (top-right) to update multiple parts at once. "
        "An editable table opens — update quantities and submit in one action."
    )

    add_heading(doc, "6.6  Location and History Views", level=2)
    add_body(doc,
        "Switch to the Location tab to see where each spare is physically stored on board. "
        "The History tab shows a full consumption and receipt log for each part."
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # 7. STORES INVENTORY
    # ─────────────────────────────────────────────
    add_heading(doc, "7.  Stores Inventory", level=1)
    add_body(doc,
        "The Stores module manages consumable inventory: general stores, lubricating oils "
        "(lubes), chemicals, and other on-board consumables. It is distinct from Spares — "
        "stores items are not component-specific and are consumed through day-to-day "
        "operations rather than specific maintenance jobs."
    )

    add_heading(doc, "7.1  Category Tabs", level=2)
    add_body(doc, "The top of the Stores screen has four category tabs:")
    add_bullet(doc, "General consumable stores items.", bold_prefix="Stores: ")
    add_bullet(doc, "Lubricating oils and greases.", bold_prefix="Lubes: ")
    add_bullet(doc, "Cleaning agents, water treatment chemicals, and similar.", bold_prefix="Chemicals: ")
    add_bullet(doc, "Miscellaneous items that do not fit the above categories.", bold_prefix="Others: ")

    img = annotate_screenshot(
        "06_stores.jpg", "07_stores_tabs.png",
        boxes=[
            (508, 88, 832, 126),   # category tabs
            (508, 140, 820, 178),  # sub-tabs
            (105, 193, 1260, 250), # table header
        ],
        labels=["Category Tabs", "Inventory/Location/History", "Table Columns"]
    )
    add_screenshot(doc, img,
        "Figure 7.1 – Stores Inventory: category tabs (Stores/Lubes/Chemicals/Others), "
        "sub-tabs, and table column headers"
    )

    add_heading(doc, "7.2  Inventory, Location, and History Views", level=2)
    add_body(doc, "Within each category, three sub-tabs are available:")
    add_bullet(doc, "Displays current ROB, minimum stock levels, and UOM (unit of measure).", bold_prefix="Inventory: ")
    add_bullet(doc, "Shows the physical storage location of each item on board.", bold_prefix="Location: ")
    add_bullet(doc, "Full audit trail of receipts, issues, and adjustments.", bold_prefix="History: ")

    add_heading(doc, "7.3  Table Columns", level=2)
    add_body(doc, "The Stores Inventory table includes:")
    add_bullet(doc, "Item Code and Item Name — unique identifier and description.")
    add_bullet(doc, "Stores Category — the classification within the selected tab.")
    add_bullet(doc, "UOM — Unit of Measure (litres, kg, units, etc.).")
    add_bullet(doc, "ROB — current Remaining on Board quantity.")
    add_bullet(doc, "Min — minimum acceptable stock level.")
    add_bullet(doc, "Stock — a status indicator (In Stock / Low Stock / Out of Stock).")
    add_bullet(doc, "Location — where the item is stored on board.")
    add_bullet(doc, "IHM — whether the item is included in the Inventory of Hazardous Materials.")
    add_bullet(doc, "Actions — edit or delete the item.")

    add_heading(doc, "7.4  Adding a New Store Item", level=2)
    add_body(doc, "To add a new consumable item to a stores category:")
    add_step(doc, 1, "Select the appropriate category tab (Stores, Lubes, Chemicals, or Others).")
    add_step(doc, 2, "Click the Add Store button (top-right).")
    add_step(doc, 3, "Fill in the item code, name, UOM, minimum stock level, and location.")
    add_step(doc, 4, "Click Save. The item is added to the inventory.")

    add_heading(doc, "7.5  Bulk Update and Export", level=2)
    add_body(doc,
        "Use Bulk Update Stores to update ROB figures for multiple items at once. "
        "Use the Export button to download the current stores list as an Excel file."
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # 8. REPORTS
    # ─────────────────────────────────────────────
    add_heading(doc, "8.  Reports", level=1)
    add_body(doc,
        "The Reports module provides access to all standard and compliance reports for the "
        "PMS. Reports can be generated for any vessel and date range, and are available for "
        "download in PDF or Excel format."
    )

    add_heading(doc, "8.1  Report Categories", level=2)
    add_body(doc, "Reports are organized into the following categories (visible in the left tree):")
    report_cats = [
        ("Maintenance & Work Orders", "Work order status, overdue, postponed, and completion reports."),
        ("Running Hours & Condition", "Running hours records and component utilization."),
        ("Inventory – Spares", "Spare parts ROB, low stock, and consumption reports."),
        ("Inventory – Stores/Lubes/Ch.", "Consumable stores and lubricant inventory reports."),
        ("IHM", "Inventory of Hazardous Materials compliance reports."),
        ("Modify PMS (Change Requests)", "Summary of submitted, approved, and rejected change requests."),
        ("Critical Equipment", "Maintenance status for critical machinery items only."),
        ("LSA/FFA Equipment", "Life-Saving Appliances and Fire Fighting Appliances reports."),
    ]
    for cat, desc in report_cats:
        add_bullet(doc, f" {desc}", bold_prefix=f"{cat}: ")

    img = annotate_screenshot(
        "07_reports.jpg", "08_reports_tree.png",
        boxes=[
            (105, 193, 408, 545),  # reports tree
            (105, 140, 540, 178),  # search and period filter
        ],
        labels=["Report Categories", "Search & Period Filter"]
    )
    add_screenshot(doc, img,
        "Figure 8.1 – Reports module: category tree (left) with report groups, "
        "search and period filter above"
    )

    add_heading(doc, "8.2  Generating a Report", level=2)
    add_body(doc, "To generate a report:")
    add_step(doc, 1, "Expand a category in the left tree by clicking the arrow beside it.")
    add_step(doc, 2, "Click the specific report name. The report configuration panel opens on the right.")
    add_step(doc, 3, "Set the Period (date range) and any additional parameters.")
    add_step(doc, 4, "Click Generate. The report is compiled in real-time.")
    add_step(doc, 5, "Click Download to save the report as a PDF or Excel file.")

    add_heading(doc, "8.3  Filters", level=2)
    add_body(doc,
        "The filter bar at the top of the Reports screen allows you to pre-filter by period "
        "and by department (All Departments or a specific HOD scope). These filters apply "
        "as default parameters when opening any report."
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # 9. MODIFY PMS
    # ─────────────────────────────────────────────
    add_heading(doc, "9.  Modify PMS", level=1)
    add_body(doc,
        "The Modify PMS module provides a governed change-management process for modifying "
        "the PMS structure. Any change to a component record, maintenance job, spare part "
        "specification, or stores item must be submitted as a formal Change Request (CR) "
        "and approved before it takes effect. This ensures traceability and compliance with "
        "class and company requirements."
    )

    add_heading(doc, "9.1  The Change Request List", level=2)
    add_body(doc,
        "The main screen shows all change requests with a category panel on the left and "
        "the requests list on the right. The status filter tabs in the header allow you to "
        "view All, Pending Approval, Approved, or Rejected change requests."
    )

    img = annotate_screenshot(
        "08_modify_pms.jpg", "09_modify_pms_overview.png",
        boxes=[
            (486, 88, 858, 126),   # status filter tabs
            (1040, 88, 1265, 126), # New Change Request button
            (105, 193, 320, 410),  # category sidebar
        ],
        labels=["Status Tabs", "New Change Request", "Category Filter"]
    )
    add_screenshot(doc, img,
        "Figure 9.1 – Modify PMS screen: status filter tabs (All/Pending Approval/Approved/Rejected), "
        "New Change Request button, and category sidebar"
    )

    add_heading(doc, "9.2  Change Request Categories", level=2)
    add_body(doc, "Change requests are grouped into four categories (left panel):")
    add_bullet(doc, "Add, remove, or edit component records in the equipment hierarchy.", bold_prefix="1. Components: ")
    add_bullet(doc, "Modify job titles, frequencies, instructions, or assigned ranks.", bold_prefix="2. Jobs: ")
    add_bullet(doc, "Update spare part specifications, part numbers, or minimum stock levels.", bold_prefix="3. Spares: ")
    add_bullet(doc, "Update stores item specifications, UOM, or minimum stock levels.", bold_prefix="4. Stores: ")

    add_heading(doc, "9.3  Submitting a Change Request", level=2)
    add_body(doc, "To create a new change request:")
    add_step(doc, 1, "Click the New Change Request button (top-right, green).")
    add_step(doc, 2, "Select the category (Components, Jobs, Spares, or Stores).")
    add_step(doc, 3, "Use the target picker to select the specific item to change.")
    add_step(doc, 4, "Describe the proposed change: fill in the new values in the Proposed fields.")
    add_step(doc, 5, "Add supporting remarks or attachments if required.")
    add_step(doc, 6, "Click Submit. The CR is sent to the superintendent for review.")

    add_heading(doc, "9.4  Reviewing and Approving a Change Request", level=2)
    add_body(doc,
        "Superintendents and authorized shore users see pending CRs in the Pending Approval tab. "
        "The review panel shows a side-by-side comparison of the Original values and the Proposed "
        "values, making it easy to evaluate the requested change."
    )
    add_bullet(doc, "The change is applied to the PMS immediately. The CR record shows the approver and timestamp.", bold_prefix="Approve: ")
    add_bullet(doc, "The CR is returned with a rejection reason. The submitter is notified.", bold_prefix="Reject: ")

    add_heading(doc, "9.5  Filters and Search", level=2)
    add_body(doc,
        "Use the search bar to find change requests by keyword. The Period date picker filters "
        "CRs by submission date. Click Clear to reset all active filters. Use the category "
        "sidebar to show only CRs of a specific type."
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # APPENDIX: ROLE-BASED ACCESS
    # ─────────────────────────────────────────────
    add_heading(doc, "Appendix A: Role-Based Access Summary", level=1)
    add_body(doc,
        "Different user roles have different levels of access across the PMS module. "
        "The table below summarises the key permissions by role."
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["Feature", "Vessel Crew", "Vessel Admin (Chief Eng.)", "Superintendent", "Client Admin"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1B3A6B')
        tcPr.append(shd)

    rows_data = [
        ("Dashboard – View", "✓", "✓", "✓", "✓"),
        ("Work Orders – Execute", "✓", "✓", "✓", "✓"),
        ("Work Orders – Approve", "–", "✓", "✓", "✓"),
        ("Running Hours – Update", "✓", "✓", "✓", "✓"),
        ("Spares ROB – Update", "✓", "✓", "✓", "✓"),
        ("Stores – Add/Edit", "–", "✓", "✓", "✓"),
        ("Modify PMS – Create CR", "–", "✓", "✓", "✓"),
        ("Modify PMS – Approve CR", "–", "–", "✓", "✓"),
        ("Admin – Configuration", "–", "–", "–", "✓"),
        ("Reports – Generate", "✓", "✓", "✓", "✓"),
    ]
    for row_data in rows_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Note: Vessel Admin = Head of Department role. "
        "The New Change Request button is hidden for Vessel Crew roles. "
        "Approval actions require Superintendent or higher privileges."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(10)

    doc.save(OUTPUT_FILE)
    print(f"✓ Document saved: {OUTPUT_FILE}")
    print(f"  Pages: ~{len(doc.sections)} sections generated")


if __name__ == "__main__":
    build_manual()
